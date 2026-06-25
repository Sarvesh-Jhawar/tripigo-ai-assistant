from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

def add_heading(doc, text, level=1):
    heading = doc.add_heading(text, level=level)
    return heading

def add_paragraph(doc, text, style=None):
    p = doc.add_paragraph(text)
    if style:
        p.style = style
    return p

doc = Document()

# Title
title = doc.add_heading('Tripigo Tales Chatbot - Complete Project Documentation', 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

add_paragraph(doc, 'Author: Sarvesh Jhawar')
add_paragraph(doc, 'Overview')
add_paragraph(doc, 'This document provides a comprehensive guide on the Tripigo Tales Chatbot project. It outlines the system architecture, detailed folder structure, backend endpoints, how to run the project locally, how to deploy it on a Linux server, and how to integrate a frontend with the backend.')

add_heading(doc, '1. System Architecture', level=1)
add_paragraph(doc, 'The Tripigo Tales Chatbot is built upon a Retrieval-Augmented Generation (RAG) architecture combining Hybrid Search with Large Language Models (LLMs). The architecture is divided into the following key components:')
add_paragraph(doc, '- FastAPI Backend: Provides a robust, asynchronous RESTful API framework routing all user requests.', style='List Bullet')
add_paragraph(doc, '- Embedding Model: Uses the HuggingFace "BAAI/bge-small-en-v1.5" sentence transformer model to convert travel data and user queries into dense vector embeddings.', style='List Bullet')
add_paragraph(doc, '- Hybrid Search Engine: Combines two retrieval methods using Reciprocal Rank Fusion (RRF):', style='List Bullet')
add_paragraph(doc, '   1. Vector Search: Powered by ChromaDB for semantic similarity matching.', style='List Number')
add_paragraph(doc, '   2. Keyword Search: Powered by BM25Okapi for exact keyword and term matching.', style='List Number')
add_paragraph(doc, '- LLM Integration: Uses the Groq API (specifically the "llama-3.1-8b-instant" model) to generate intelligent, well-formatted markdown responses strictly based on the retrieved context (anti-hallucination enforced).', style='List Bullet')

add_heading(doc, '2. Detailed Project File Structure', level=1)
add_paragraph(doc, 'Here is an exhaustive breakdown of every file and folder in the project and its specific purpose:')

add_heading(doc, 'Root Directory', level=2)
add_paragraph(doc, '- app.py: The main entry point for the FastAPI backend application. It initializes the app, configures CORS middleware to allow cross-origin requests, and includes the routers for the /chat, /ingest, and /health endpoints.', style='List Bullet')
add_paragraph(doc, '- requirements.txt: Lists all Python dependencies (FastAPI, uvicorn, chromadb, sentence-transformers, groq, rank-bm25, etc.) required to run the project.', style='List Bullet')
add_paragraph(doc, '- .env: Stores environment variables securely, such as the GROQ_API_KEY.', style='List Bullet')
add_paragraph(doc, '- generate_docs.py: A Python script (using python-docx) to dynamically generate this DOCX documentation.', style='List Bullet')

add_heading(doc, 'routes/ Directory', level=2)
add_paragraph(doc, 'Handles all API routing and HTTP request/response validation logic.')
add_paragraph(doc, '- routes/chat.py: Defines the `/chat` POST endpoint. It takes the user\'s message, triggers the hybrid search to retrieve context, passes the context to the LLM, and formats the response (appending sources).', style='List Bullet')
add_paragraph(doc, '- routes/ingest.py: Defines the `/ingest` POST endpoint. It reads data from `tripigo_chunks.json`, creates embeddings, stores them in ChromaDB, and initializes the BM25 index.', style='List Bullet')
add_paragraph(doc, '- routes/health.py: Defines the `/health` GET endpoint for basic health checks.', style='List Bullet')

add_heading(doc, 'services/ Directory', level=2)
add_paragraph(doc, 'Houses the core business logic, RAG system, and external API integrations.')
add_paragraph(doc, '- services/llm_service.py: Contains the Groq client setup and the `generate_answer` function. It holds the strict System Prompt that enforces formatting rules and prevents hallucinations. It uses the `llama-3.1-8b-instant` model.', style='List Bullet')
add_paragraph(doc, '- services/rag_service.py: The core search engine logic. It manages ChromaDB collections, maintains the BM25 index, and implements the `hybrid_search` algorithm to fuse vector and keyword results.', style='List Bullet')
add_paragraph(doc, '- services/vector_service.py: Initializes the SentenceTransformer model (`BAAI/bge-small-en-v1.5`) and provides utility functions to embed text strings and documents into vectors.', style='List Bullet')

add_heading(doc, 'Data & Storage Directories', level=2)
add_paragraph(doc, '- data/tripigo_chunks.json: Raw JSON file containing all travel packages, policies, prices, and keywords. This is the source of truth for the chatbot.', style='List Bullet')
add_paragraph(doc, '- chroma_db/: The persistent local directory where ChromaDB stores vectorized embeddings. It allows the chatbot to instantly query data without re-embedding it every time.', style='List Bullet')

add_heading(doc, '3. API Keys & Configuration', level=1)
add_paragraph(doc, 'To make the application work, you need an API key.')
add_paragraph(doc, 'Instructions: visit this website and in developers drop down there would be generate api kkey take that key and put in .env place holder.')

add_heading(doc, '4. Backend Endpoints', level=1)
add_paragraph(doc, 'The backend is built with FastAPI. Below are the endpoints, their purposes, and how to request them.')

add_heading(doc, 'Endpoint: /chat (POST)', level=2)
add_paragraph(doc, 'Purpose: The main endpoint for communicating with the AI.')
add_paragraph(doc, 'Request Body (JSON):', style='List Bullet')
add_paragraph(doc, '{"message": "I am looking for a honeymoon package under 20000"}')
add_paragraph(doc, 'Response Structure (JSON):', style='List Bullet')
add_paragraph(doc, '{\n  "answer": "...",\n  "sources": ["..."]\n}')
add_paragraph(doc, 'Description: The "answer" field will contain a beautifully formatted Markdown string with the AI\'s response, complete with headings, bold text, and bullet points.')

add_heading(doc, 'Endpoint: /ingest (POST)', level=2)
add_paragraph(doc, 'Purpose: Ingests data from the local data files (e.g., data/tripigo_chunks.json) into the ChromaDB vector database and BM25 index.')
add_paragraph(doc, 'When to use: Call this endpoint whenever you update your travel packages or policies in the data JSON files to sync the AI\'s knowledge.')
add_paragraph(doc, 'Request Body: None')

add_heading(doc, 'Endpoint: /health (GET)', level=2)
add_paragraph(doc, 'Purpose: A simple health check to ensure the API is running correctly.')
add_paragraph(doc, 'Returns: {"status": "healthy"}')

add_heading(doc, '5. How to Run on Localhost', level=1)
add_paragraph(doc, 'Follow these steps to run the backend locally:')
add_paragraph(doc, '1. Ensure you have Python 3.9+ installed.', style='List Number')
add_paragraph(doc, '2. Create a virtual environment: `python -m venv venv`', style='List Number')
add_paragraph(doc, '3. Activate the virtual environment:', style='List Number')
add_paragraph(doc, '   - Windows: `venv\\Scripts\\activate`')
add_paragraph(doc, '   - Mac/Linux: `source venv/bin/activate`')
add_paragraph(doc, '4. Install the requirements: `pip install -r requirements.txt`', style='List Number')
add_paragraph(doc, '5. Set up the `.env` file with your API key as described in Section 3.', style='List Number')
add_paragraph(doc, '6. Start the FastAPI server: `uvicorn app:app`', style='List Number')
add_paragraph(doc, 'The backend will now be running on http://localhost:8000.')

add_heading(doc, '6. Customizing Answer Length & Formatting', level=1)
add_paragraph(doc, 'You can easily increase or decrease the length of the AI\'s responses by modifying the `SYSTEM_PROMPT` inside `services/llm_service.py`.')
add_paragraph(doc, 'Example: To make the bot provide much shorter, concise answers, you can add this instruction to the SYSTEM_PROMPT:')
add_paragraph(doc, '"Keep your answers brief, maximum 3-4 sentences. Do not provide unnecessary details unless explicitly asked."')
add_paragraph(doc, 'Example: To make the bot provide longer, highly detailed answers, you can add this instruction:')
add_paragraph(doc, '"Provide exhaustive, highly detailed explanations. Expand on the destination\'s culture, best times to visit, and include a daily itinerary breakdown if possible."')

add_heading(doc, '7. How to Implement Frontend for this Backend (Example)', level=1)
add_paragraph(doc, 'The `frontend-example` folder contains a demo React app. To integrate a frontend with this backend:')
add_paragraph(doc, '1. Navigate to the frontend directory: `cd frontend-example`', style='List Number')
add_paragraph(doc, '2. Install Node.js dependencies: `npm install`', style='List Number')
add_paragraph(doc, '3. Run the frontend development server: `npm run dev`', style='List Number')
add_paragraph(doc, 'Frontend Integration Logic:', style='List Number')
add_paragraph(doc, 'Your frontend should send a POST request to `http://localhost:8000/chat` with the user\'s message.')
add_paragraph(doc, 'Because the backend returns Markdown, you must use a Markdown renderer in your frontend (e.g., `react-markdown` and `remark-gfm`).')
add_paragraph(doc, 'Example API Call (JavaScript):')
add_paragraph(doc, 'const response = await fetch("http://localhost:8000/chat", {\n  method: "POST",\n  headers: { "Content-Type": "application/json" },\n  body: JSON.stringify({ message: "Hello" })\n});\nconst data = await response.json();\nconsole.log(data.answer); // This is the markdown text\n')
add_paragraph(doc, 'For anti-hallucination, if the bot says "I couldn\'t find that information in Tripigo\'s travel database.", you can trigger a "Contact Human Support" button in your UI.')

add_heading(doc, '8. How to Deploy this Project on a Linux Server', level=1)
add_paragraph(doc, 'To deploy this project to a production Linux server (e.g., AWS EC2, DigitalOcean Droplet, Ubuntu):')
add_paragraph(doc, '1. Update your server packages: `sudo apt update && sudo apt upgrade -y`', style='List Number')
add_paragraph(doc, '2. Install Python, pip, and nginx: `sudo apt install python3-pip python3-venv nginx -y`', style='List Number')
add_paragraph(doc, '3. Clone the repository and navigate into the folder.', style='List Number')
add_paragraph(doc, '4. Create and activate a virtual environment, then install dependencies (`pip install -r requirements.txt`).', style='List Number')
add_paragraph(doc, '5. Create the `.env` file and insert your API keys.', style='List Number')
add_paragraph(doc, '6. Use Gunicorn with Uvicorn workers to run the app in the background:', style='List Number')
add_paragraph(doc, '   `gunicorn -k uvicorn.workers.UvicornWorker app:app -b 127.0.0.1:8000`')
add_paragraph(doc, '7. (Optional but Recommended) Set up a systemd service file (`/etc/systemd/system/tripigo.service`) to keep the FastAPI app running automatically on reboots.', style='List Number')
add_paragraph(doc, '8. Configure Nginx as a reverse proxy:', style='List Number')
add_paragraph(doc, '   Create a file in `/etc/nginx/sites-available/tripigo` routing traffic from port 80 to `http://127.0.0.1:8000`.')
add_paragraph(doc, '9. Enable the site and restart Nginx: `sudo systemctl restart nginx`.', style='List Number')
add_paragraph(doc, '10. Your API is now live! Anyone can integrate with it by sending requests to your server\'s IP or domain name.', style='List Number')

doc.save('Tripigo_Complete_Documentation_v2.docx')
print("Documentation generated successfully. Saved as Tripigo_Complete_Documentation_v2.docx")
